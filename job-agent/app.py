import io
import re

from flask import Flask, jsonify, request, send_file, send_from_directory

from orchestrator import run_pipeline
from utils.openrouter import OpenRouterError
from utils.pdf_export import build_resume_pdf_bytes
from utils.pdf_parser import PdfParserError, extract_pdf_text
from utils.resume_blocks import parse_resume_markdown

app = Flask(__name__, static_folder="frontend", static_url_path="")


def _safe_slug(value: str) -> str:
    text = (value or "").strip().lower()
    text = re.sub(r"[^a-z0-9]+", "-", text)
    text = re.sub(r"(^-+)|(-+$)", "", text)
    return text[:60] or "output"


def _markdown_to_plain_text(md: str) -> str:
    text = md or ""
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\-\s+", "• ", text, flags=re.MULTILINE)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()


def _make_docx_bytes(title: str, body_text: str, is_resume: bool = False) -> bytes:
    from docx import Document

    doc = Document()
    if is_resume:
        blocks = parse_resume_markdown(body_text)
        for block in blocks:
            btype = block["type"]
            value = block["value"]
            if btype == "name":
                doc.add_heading(value, level=1)
            elif btype == "section":
                doc.add_heading(value, level=2)
            elif btype == "subsection":
                doc.add_paragraph(value)
            elif btype == "bullet":
                doc.add_paragraph(value, style="List Bullet")
            elif btype == "text":
                doc.add_paragraph(value)
            else:
                doc.add_paragraph("")
    else:
        if title:
            doc.add_heading(title, level=1)
        for line in (body_text or "").splitlines():
            if not line.strip():
                doc.add_paragraph("")
            else:
                doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_pdf_bytes_cover_letter(title: str, body_text: str) -> bytes:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import inch
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=LETTER)
    width, height = LETTER
    left = 0.75 * inch
    top = height - 0.75 * inch
    y = top
    c.setTitle(title or "document")
    max_width = width - (2 * left)

    def ensure_space(y_pos: float, min_space: float = 0.75 * inch) -> float:
        if y_pos < min_space:
            c.showPage()
            return top
        return y_pos

    def draw_wrapped_line(line: str, y_pos: float, font_name: str, font_size: float) -> float:
        c.setFont(font_name, font_size)
        words = (line or "").split()
        if not words:
            return y_pos - 0.14 * inch
        current = ""
        for w in words:
            candidate = (current + " " + w).strip()
            if c.stringWidth(candidate, font_name, font_size) <= max_width:
                current = candidate
            else:
                c.drawString(left, y_pos, current)
                y_pos -= 0.16 * inch
                current = w
                y_pos = ensure_space(y_pos)
        if current:
            c.drawString(left, y_pos, current)
            y_pos -= 0.16 * inch
        return y_pos

    c.setFont("Helvetica-Bold", 14)
    if title:
        c.drawString(left, y, title)
        y -= 0.35 * inch
    for raw_line in (body_text or "").splitlines():
        y = ensure_space(y)
        if not raw_line.strip():
            y -= 0.16 * inch
            continue
        y = draw_wrapped_line(raw_line, y, "Helvetica", 10.5)
    c.save()
    return buf.getvalue()


@app.route("/", methods=["GET"])
def index():
    return send_from_directory(app.static_folder, "index.html")


@app.route("/analyze", methods=["POST"])
def analyze():
    company = (request.form.get("company") or "").strip()
    jd_text = (request.form.get("jd_text") or "").strip()
    resume_text_input = (request.form.get("resume_text") or "").strip()
    resume_preset = (request.form.get("resume_preset") or "one_page").strip()
    fact_check_raw = request.form.get("fact_check", "true")

    if resume_preset not in ("one_page", "two_page"):
        resume_preset = "one_page"

    if not company:
        return jsonify({"error": "Missing required field: company"}), 400
    if not jd_text:
        return jsonify({"error": "Missing required field: jd_text"}), 400

    resume_text = resume_text_input
    if "resume_pdf" in request.files and request.files["resume_pdf"].filename:
        try:
            resume_text = extract_pdf_text(request.files["resume_pdf"].read())
        except PdfParserError as exc:
            return jsonify({"error": str(exc)}), 400

    if not resume_text:
        return jsonify({"error": "Provide either resume_text or resume_pdf."}), 400

    try:
        result = run_pipeline(
            resume_text,
            jd_text,
            company,
            options={
                "resume_preset": resume_preset,
                "fact_check": fact_check_raw,
            },
        )
    except OpenRouterError as exc:
        return jsonify({"error": str(exc)}), 502
    except Exception as exc:
        return jsonify({"error": f"Unexpected server error: {exc}"}), 500

    return jsonify(result)


@app.route("/export", methods=["POST"])
def export():
    doc_type = (request.form.get("type") or "").strip().lower()
    fmt = (request.form.get("format") or "").strip().lower()
    company = (request.form.get("company") or "").strip()
    content = (request.form.get("content") or "").strip()

    if doc_type not in {"resume", "cover_letter"}:
        return jsonify({"error": "Invalid type. Use resume or cover_letter."}), 400
    if fmt not in {"pdf", "docx"}:
        return jsonify({"error": "Invalid format. Use pdf or docx."}), 400
    if not content:
        return jsonify({"error": "Missing content to export."}), 400

    slug = _safe_slug(company)
    title = "Resume" if doc_type == "resume" else "Cover Letter"
    body_text = content

    if fmt == "docx":
        data = _make_docx_bytes(title, body_text, is_resume=(doc_type == "resume"))
        filename = f"{slug}-{doc_type}.docx"
        return send_file(
            io.BytesIO(data),
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    if doc_type == "resume":
        data = build_resume_pdf_bytes(body_text, title=title)
    else:
        data = _make_pdf_bytes_cover_letter(title, _markdown_to_plain_text(body_text))

    filename = f"{slug}-{doc_type}.pdf"
    return send_file(
        io.BytesIO(data),
        as_attachment=True,
        download_name=filename,
        mimetype="application/pdf",
    )


if __name__ == "__main__":
    app.run(debug=True)
